#!/usr/bin/env python3
"""
Generate HPT-QRC Poster DOCX for EPFL Quantum Hackathon 2026
"""

import io
import math
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
import matplotlib.patheffects as pe
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json
import os

ASSETS = '/Users/umut/Desktop/EPFL_ANTI/qedi_website/assets'

# ─── COLOURS ──────────────────────────────────────────────────────────────────
EPFL_RED   = RGBColor(0xE3, 0x00, 0x0F)
DARK_SLATE = RGBColor(0x1E, 0x29, 0x3B)
LIGHT_BG   = RGBColor(0xF8, 0xFA, 0xFC)
ACCENT_GRN = RGBColor(0x2E, 0xCC, 0x71)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
MID_GREY   = RGBColor(0x64, 0x74, 0x8B)


# ─── HELPER: set paragraph shading ───────────────────────────────────────────
def set_para_shading(para, fill_hex: str):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)


def set_cell_bg(cell, fill_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    """Set borders on a cell. kwargs: top/bottom/left/right = (size, color, val)"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        if side in kwargs:
            sz, color, val = kwargs[side]
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val)
            el.set(qn('w:sz'), str(sz))
            el.set(qn('w:color'), color)
            tcBorders.append(el)
    tcPr.append(tcBorders)


def add_run(para, text, bold=False, italic=False, size=11, color=None, font='Calibri'):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font
    if color:
        run.font.color.rgb = color
    return run


def add_heading(doc, text, level=1, color=DARK_SLATE, size=None):
    sizes = {1: 20, 2: 15, 3: 13, 4: 11}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size or sizes.get(level, 12))
    run.font.name = 'Calibri'
    run.font.color.rgb = color
    return p


def add_divider(doc, color_hex='E3000F'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def img_buf(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    buf.seek(0)
    plt.close(fig)
    return buf


# ─── PIPELINE DIAGRAM ─────────────────────────────────────────────────────────
def make_pipeline_image():
    fig, ax = plt.subplots(figsize=(18, 4.5))
    ax.set_xlim(0, 18)
    ax.set_ylim(0, 4.5)
    ax.axis('off')
    fig.patch.set_facecolor('#0F172A')

    steps = [
        # (x_center, width, label_lines, color, text_color)
        (1.1,  1.8, ['Raw Market', 'Surface', '(224D)'],        '#1E3A5F', 'white'),
        (3.4,  2.0, ['Standard Scaler', '+', 'PCA → 5D'],       '#1E3A5F', 'white'),
        (5.7,  2.0, ['Rolling 5-Day', 'Window', '(1 × 25)'],    '#1E3A5F', 'white'),
        (9.2,  5.6, ['Hybrid Photonic Temporal QRC',
                     '5 Input Modes  |  3 Memory Modes',
                     '3 Seeds × 3 Virtual Nodes = 9 Circuits'], '#7C1019', 'white'),
        (13.0, 2.4, ['LexGrouping (90Q)', '+', 'Window (25C)'], '#1E3A5F', '#93C5FD'),
        (15.4, 2.0, ['L2 Ridge', 'Regression',
                     'Inv PCA → Inv Scaler'],                    '#1E3A5F', 'white'),
        (17.2, 1.6, ['Day T+1', 'Forecast', '(224D)'],          '#14532D', '#6EE7B7'),
    ]

    box_h = 2.2
    y0 = (4.5 - box_h) / 2

    arrow_kw = dict(arrowstyle='->', color='#94A3B8', lw=1.8,
                    mutation_scale=14, connectionstyle='arc3,rad=0')

    prev_right = None
    for (cx, w, lines, fc, tc) in steps:
        x0 = cx - w / 2
        rect = FancyBboxPatch((x0, y0), w, box_h,
                              boxstyle='round,pad=0.12',
                              facecolor=fc, edgecolor='#334155', linewidth=1.2, zorder=3)
        ax.add_patch(rect)

        # Label lines
        n = len(lines)
        for i, line in enumerate(lines):
            yi = y0 + box_h * (n - i) / (n + 1)
            fs = 9.5 if i == 0 and (cx == 9.2) else 8.2
            fw = 'bold' if i == 0 else 'normal'
            ax.text(cx, yi, line, ha='center', va='center',
                    color=tc, fontsize=fs, fontweight=fw, zorder=4)

        if prev_right is not None:
            ax.annotate('', xy=(x0 + 0.05, y0 + box_h / 2),
                        xytext=(prev_right - 0.05, y0 + box_h / 2),
                        arrowprops=arrow_kw, zorder=5)
        prev_right = x0 + w

    # Decorative "photon" dots inside the QRC block
    qrc_cx, qrc_w = 9.2, 5.6
    np.random.seed(42)
    for _ in range(18):
        px = qrc_cx - qrc_w / 2 + 0.3 + np.random.rand() * (qrc_w - 0.6)
        py = y0 + 0.25 + np.random.rand() * (box_h - 0.5)
        ax.plot(px, py, 'o', color='#FF7675', ms=3.5, alpha=0.7, zorder=5)

    ax.set_title('HPT-QRC Pipeline Architecture',
                 color='white', fontsize=13, fontweight='bold', pad=10)
    return img_buf(fig)


# ─── MODEL COMPARISON BAR CHART ───────────────────────────────────────────────
def make_model_chart():
    models  = ['QSVR', 'Hybrid\nQNN', 'LSTM', 'Photonic\nLinear QRC',
               'Hybrid Photonic\nLinear QRC', 'HPT-QRC\n(Ours)']
    rmse    = [0.0233,  0.0083,  0.0073,  0.0065,  0.0028,  0.0021]
    colors  = ['#64748B', '#64748B', '#64748B', '#64748B', '#64748B', '#E3000F']

    fig, ax = plt.subplots(figsize=(10, 4))
    fig.patch.set_facecolor('#F8FAFC')
    ax.set_facecolor('#F8FAFC')

    bars = ax.bar(models, rmse, color=colors, width=0.6, zorder=3, edgecolor='white', linewidth=0.5)
    for bar, val in zip(bars, rmse):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.0003,
                f'{val:.4f}', ha='center', va='bottom', fontsize=9, fontweight='bold',
                color='#1E293B')

    ax.set_ylabel('RMSE (lower is better)', fontsize=10, color='#1E293B')
    ax.set_title('Model Performance Comparison', fontsize=12, fontweight='bold', color='#1E293B')
    ax.spines[['top', 'right']].set_visible(False)
    ax.tick_params(colors='#1E293B', labelsize=9)
    ax.yaxis.grid(True, alpha=0.4, zorder=0)
    ax.set_ylim(0, 0.027)

    # Annotate champion
    ax.annotate('Champion Model', xy=(5, 0.0021), xytext=(4.2, 0.010),
                arrowprops=dict(arrowstyle='->', color='#E3000F', lw=1.5),
                color='#E3000F', fontsize=9, fontweight='bold')

    fig.tight_layout()
    return img_buf(fig)


# ─── METRICS TABLE IMAGE ──────────────────────────────────────────────────────
def make_metrics_table():
    data = [
        ['Model',                   'MSE',      'RMSE',  'MAE'],
        ['Classical LSTM',          '5.26e-05', '0.0073', '0.0052'],
        ['QSVR',                    '5.43e-04', '0.0233', '0.0187'],
        ['Hybrid QNN',              '6.85e-05', '0.0083', '0.0059'],
        ['Photonic Linear QRC',     '4.22e-05', '0.0065', '0.0045'],
        ['Hybrid Photonic Lin. QRC','7.70e-06', '0.0028', '0.0021'],
        ['HPT-QRC (Ours) 🥇',       '7.58e-06', '0.0021', '0.00195'],
    ]

    fig, ax = plt.subplots(figsize=(9, 3.2))
    fig.patch.set_facecolor('white')
    ax.axis('off')

    col_colors = [['#1E293B'] * 4]
    row_colors = [['#F1F5F9'] * 4 if i % 2 == 0 else ['white'] * 4
                  for i in range(len(data) - 1)]
    row_colors[-1] = ['#FEF2F2'] * 4   # champion row

    table = ax.table(
        cellText=[row for row in data[1:]],
        colLabels=data[0],
        loc='center',
        cellLoc='center',
        colColours=['#1E293B'] * 4,
        cellColours=row_colors
    )
    table.auto_set_font_size(False)
    table.set_fontsize(10)
    table.scale(1, 1.6)

    for (row, col), cell in table.get_celld().items():
        cell.set_edgecolor('#CBD5E1')
        cell.set_linewidth(0.5)
        if row == 0:
            cell.get_text().set_color('white')
            cell.get_text().set_fontweight('bold')
        elif row == len(data) - 1:
            cell.get_text().set_color('#B91C1C')
            cell.get_text().set_fontweight('bold')

    ax.set_title('Full Performance Leaderboard', fontsize=11, fontweight='bold',
                 color='#1E293B', pad=8)
    fig.tight_layout()
    return img_buf(fig)


# ─── BUILD DOCUMENT ───────────────────────────────────────────────────────────
def build_docx():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    # ── TITLE BANNER ──────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_shading(title_p, '1E293B')
    title_p.paragraph_format.space_before = Pt(6)
    title_p.paragraph_format.space_after  = Pt(0)
    add_run(title_p, 'Hybrid Photonic Temporal QRC  (HPT-QRC)',
            bold=True, size=22, color=WHITE, font='Calibri')

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_shading(subtitle_p, '1E293B')
    subtitle_p.paragraph_format.space_before = Pt(0)
    subtitle_p.paragraph_format.space_after  = Pt(0)
    add_run(subtitle_p, 'Swaption Volatility Surface Forecaster',
            bold=False, size=15, color=RGBColor(0x93, 0xC5, 0xFD), font='Calibri')

    event_p = doc.add_paragraph()
    event_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_shading(event_p, '1E293B')
    event_p.paragraph_format.space_before = Pt(0)
    event_p.paragraph_format.space_after  = Pt(6)
    add_run(event_p, 'EPFL Quantum Hackathon 2026  ·  Quandela Challenge  ·  Team Qedi',
            size=10, color=RGBColor(0x94, 0xA3, 0xB8), font='Calibri')

    # Team & uni
    team_p = doc.add_paragraph()
    team_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_shading(team_p, 'E3000F')
    team_p.paragraph_format.space_before = Pt(0)
    team_p.paragraph_format.space_after  = Pt(4)
    add_run(team_p,
            'Eren Aslan  ·  Hüseyin Umut Işık  ·  Arda Kara  ·  Mehmet Alp Özaydın   |   METU & Bilkent University',
            bold=True, size=10.5, color=WHITE, font='Calibri')

    doc.add_paragraph()  # spacer

    # ── TWO-COLUMN LAYOUT: Problem + Theory ───────────────────────────────────
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.columns[0].width = Cm(8.8)
    tbl.columns[1].width = Cm(8.8)

    # ---- LEFT: The Problem ----
    cell_L = tbl.cell(0, 0)
    set_cell_bg(cell_L, 'F8FAFC')
    set_cell_border(cell_L, right=(4, 'E3000F', 'single'))

    prob_head = cell_L.paragraphs[0]
    prob_head.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(prob_head, '01.  THE CHALLENGE', bold=True, size=13,
            color=EPFL_RED, font='Calibri')

    def cp(cell, text='', bold=False, italic=False, size=10.5, color=DARK_SLATE,
           align=WD_ALIGN_PARAGRAPH.LEFT, space_before=2, space_after=2):
        p = cell.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        if text:
            add_run(p, text, bold=bold, italic=italic, size=size, color=color)
        return p

    cp(cell_L, '')
    cp(cell_L,
       'Swaption volatility surfaces encode the market\'s expectation of interest-rate '
       'volatility across all tenor–maturity combinations. Each surface is a '
       '224-dimensional snapshot that changes daily.', size=10)

    cp(cell_L, 'Task 1 — Predictive Forecasting', bold=True, size=10.5,
       color=EPFL_RED, space_before=6)
    cp(cell_L,
       'Predict the next H-day swaption volatility surface topology from historical '
       'market data.  Classical models (LSTM, regression) struggle with the '
       'high-dimensional, non-linear structure of these surfaces.', size=10)

    cp(cell_L, 'Task 2 — Data Imputation', bold=True, size=10.5,
       color=EPFL_RED, space_before=6)
    cp(cell_L,
       'Reconstruct missing option price nodes in the 2-D surface grid.  '
       'Real market feeds often contain gaps, making robust imputation critical '
       'for downstream risk models.', size=10)

    cp(cell_L, 'Why Quantum?', bold=True, size=10.5, color=EPFL_RED, space_before=6)
    cp(cell_L,
       'Quantum reservoirs naturally encode high-dimensional, non-linear '
       'feature spaces through interference and entanglement — properties that '
       'classical architectures must approximate at great computational cost.', size=10)

    # ---- RIGHT: Theory ----
    cell_R = tbl.cell(0, 1)
    set_cell_bg(cell_R, 'F0F4FF')

    theory_head = cell_R.paragraphs[0]
    theory_head.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(theory_head, '02.  THEORY & BACKGROUND', bold=True, size=13,
            color=EPFL_RED, font='Calibri')

    cp(cell_R, '')
    cp(cell_R, 'Quantum Reservoir Computing (QRC)', bold=True, size=10.5,
       color=DARK_SLATE, space_before=4)
    cp(cell_R,
       'QRC exploits a fixed, high-dimensional quantum system as a '
       'non-linear dynamical reservoir.  A simple linear readout layer '
       'trained on the reservoir state is enough to solve complex '
       'time-series tasks — no gradient-based training of the quantum '
       'circuit is needed.', size=10)

    cp(cell_R, 'Photonic Advantage via MerLin', bold=True, size=10.5,
       color=DARK_SLATE, space_before=6)
    cp(cell_R,
       'We implement the reservoir on Quandela\'s MerLin photonic framework. '
       'Linear-optical circuits with beamsplitters and phase-shifters create '
       'rich Fock-space interference.  Photons do not decohere as quickly as '
       'qubits, making photonic QRC a near-term hardware candidate.', size=10)

    cp(cell_R, 'Inspiration: Li et al. (2024)', bold=True, size=10.5,
       color=DARK_SLATE, space_before=6)
    cp(cell_R,
       '"QRC for Realized Volatility Forecasting" introduced a qubit-based '
       'Hamiltonian-evolution reservoir for financial time series.  '
       'We adapted and extended this to a purely photonic setting, adding '
       'dedicated memory modes, virtual nodes, and ensemble LexGrouping.', size=10)

    cp(cell_R, 'PCA Dimensionality Reduction', bold=True, size=10.5,
       color=DARK_SLATE, space_before=6)
    cp(cell_R,
       'The 224D surface is compressed to 5 principal components (capturing '
       '>95 % of variance) before being fed into the reservoir, making the '
       'computation tractable while preserving the surface structure.', size=10)

    doc.add_paragraph()

    # ── SECTION 3: PIPELINE ────────────────────────────────────────────────────
    pipe_head = add_heading(doc, '03.  HPT-QRC PIPELINE ARCHITECTURE', 2, color=DARK_SLATE, size=14)
    pipe_head.paragraph_format.space_before = Pt(4)
    add_divider(doc)

    p_desc = doc.add_paragraph()
    add_run(p_desc,
            'End-to-end data flow from raw 224-dimensional swaption surface to next-day forecast:',
            size=10, color=MID_GREY)
    p_desc.paragraph_format.space_after = Pt(4)

    pipeline_buf = make_pipeline_image()
    doc.add_picture(pipeline_buf, width=Inches(6.8))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(cap, 'Figure 1.  Raw (224D) → Scaler+PCA(5D) → Rolling Window(1×25) → '
            'HPT-QRC(9 circuits) → LexGrouping(90Q)+Classic(25C) → Ridge → Forecast(224D)',
            italic=True, size=8.5, color=MID_GREY)

    doc.add_paragraph()

    # ── SECTION 4: KEY INNOVATIONS ────────────────────────────────────────────
    inno_head = add_heading(doc, '04.  KEY INNOVATIONS', 2, color=DARK_SLATE, size=14)
    add_divider(doc)

    inno_table = doc.add_table(rows=2, cols=2)
    inno_table.style = 'Table Grid'
    inno_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    innovations = [
        ('01  Dedicated Memory Modes',
         'Uses 5 input modes + 3 unencoded memory modes. '
         'Memory modes continuously accumulate historical context through serial '
         'phase mixing across the 5-step temporal window — giving the reservoir '
         'genuine short-term memory without additional parameters.'),
        ('02  Virtual Nodes',
         'The reservoir is sampled at multiple post-processing depths (V1, V2, V3). '
         'This emulates chronological sub-interval measurements, massively expanding '
         'the temporal feature space from 8 physical modes to 90 quantum features '
         'without increasing photon count.'),
        ('03  Ensemble LexGrouping',
         '3 random seeds × 3 virtual depths = 9 independent circuits. '
         'Raw Fock-state probabilities are compressed via LexGrouping, '
         'making the Hilbert-space output tractable. '
         'Ensembling dramatically reduces variance and outperforms single-circuit QRC.'),
        ('04  Hybrid Readout',
         '90 quantum features (LexGrouped probabilities) are concatenated with '
         '25 classical window features.  An L2-regularised Ridge regression then '
         'projects these 115 features onto the 5 PCA coordinates, which are '
         'inverse-transformed to recover the full 224D forecast surface.'),
    ]

    cells = [inno_table.cell(r, c) for r in range(2) for c in range(2)]
    bg_colors = ['F0F4FF', 'FFF0F0', 'F0FFF4', 'FFFFF0']

    for i, (title, body) in enumerate(innovations):
        c = cells[i]
        set_cell_bg(c, bg_colors[i])
        p = c.paragraphs[0]
        add_run(p, title, bold=True, size=11, color=EPFL_RED if i % 2 == 0 else DARK_SLATE)
        p.paragraph_format.space_after = Pt(3)
        bp = c.add_paragraph()
        add_run(bp, body, size=10, color=DARK_SLATE)
        bp.paragraph_format.space_before = Pt(2)

    doc.add_paragraph()

    # ── SECTION 5: PCA VARIANCE PLOT ─────────────────────────────────────────
    arch_head = add_heading(doc, '05.  DIMENSIONALITY REDUCTION', 2, color=DARK_SLATE, size=14)
    add_divider(doc)

    pca_img = os.path.join(ASSETS, 'pca_variance.jpeg')
    doc.add_picture(pca_img, width=Inches(6.0))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap2 = doc.add_paragraph()
    cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(cap2, 'Figure 2.  PCA Cumulative Explained Variance — 5 components capture >95% of surface variance.',
            italic=True, size=8.5, color=MID_GREY)

    doc.add_paragraph()

    # ── SECTION 6: RESULTS ────────────────────────────────────────────────────
    res_head = add_heading(doc, '06.  RESULTS', 2, color=DARK_SLATE, size=14)
    add_divider(doc)

    p_res = doc.add_paragraph()
    add_run(p_res,
            'HPT-QRC achieves the lowest RMSE across all baselines, outperforming LSTM, '
            'QSVR, Hybrid QNN, and even pure photonic linear QRC by a wide margin.',
            size=10.5, color=DARK_SLATE)
    p_res.paragraph_format.space_after = Pt(6)

    # Bar chart
    chart_buf = make_model_chart()
    doc.add_picture(chart_buf, width=Inches(6.2))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap3 = doc.add_paragraph()
    cap3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(cap3, 'Figure 3.  RMSE comparison across all model architectures.  HPT-QRC (red) achieves 0.0021.',
            italic=True, size=8.5, color=MID_GREY)

    doc.add_paragraph()

    # Metrics table image
    metrics_buf = make_metrics_table()
    doc.add_picture(metrics_buf, width=Inches(6.2))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap4 = doc.add_paragraph()
    cap4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(cap4, 'Table 1.  Full performance leaderboard (MSE, RMSE, MAE) across all architectures.',
            italic=True, size=8.5, color=MID_GREY)

    doc.add_paragraph()

    # Walk-forward error image
    add_heading(doc, '6a.  Temporal Walk-Forward Evaluation (6-Day Horizon)', 3,
                color=DARK_SLATE, size=12)
    wf_img = os.path.join(ASSETS, 'base_model_walk_forward_errors.jpg')
    doc.add_picture(wf_img, width=Inches(6.5))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap5 = doc.add_paragraph()
    cap5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(cap5, 'Figure 4.  6-Day walk-forward error evaluation — HPT-QRC maintains low error across all horizons.',
            italic=True, size=8.5, color=MID_GREY)

    doc.add_paragraph()

    # Configuration experiments
    add_heading(doc, '6b.  Hyperparameter Configuration Impact', 3,
                color=DARK_SLATE, size=12)
    cfg_img = os.path.join(ASSETS, 'configuration_experiments.jpg')
    doc.add_picture(cfg_img, width=Inches(6.5))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap6 = doc.add_paragraph()
    cap6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(cap6, 'Figure 5.  Hyperparameter ablation study — seeds, virtual nodes, and modes swept.',
            italic=True, size=8.5, color=MID_GREY)

    doc.add_paragraph()

    # Task 2 imputation
    add_heading(doc, '6c.  Task 2: Data Imputation Reconstruction', 3,
                color=DARK_SLATE, size=12)
    t2_img = os.path.join(ASSETS, 'task_2_arda.jpeg')
    doc.add_picture(t2_img, width=Inches(6.5))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap7 = doc.add_paragraph()
    cap7.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(cap7, 'Figure 6.  Data imputation reconstruction results — missing surface nodes recovered accurately.',
            italic=True, size=8.5, color=MID_GREY)

    doc.add_paragraph()

    # Champion price prediction
    add_heading(doc, '6d.  Champion Model: Price Trajectory Tracking', 3,
                color=DARK_SLATE, size=12)
    champ_img = os.path.join(ASSETS, 'Champion_Hybrid_QNN_Price_Prediction.png')
    doc.add_picture(champ_img, width=Inches(6.5))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap8 = doc.add_paragraph()
    cap8.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(cap8, 'Figure 7.  HPT-QRC swaption price trajectory tracking vs. actual market data.',
            italic=True, size=8.5, color=MID_GREY)

    doc.add_paragraph()

    # ── SECTION 7: CONCLUSIONS ─────────────────────────────────────────────────
    conc_head = add_heading(doc, '07.  CONCLUSIONS', 2, color=DARK_SLATE, size=14)
    add_divider(doc)

    conclusions = [
        ('State-of-the-art accuracy',
         'HPT-QRC achieves RMSE = 0.0021, beating LSTM (0.0073), QSVR (0.0233), '
         'and all intermediate photonic variants.'),
        ('Novel temporal memory architecture',
         'Dedicated memory modes + virtual-node sampling give the photonic '
         'reservoir genuine short-term memory at zero extra hardware cost.'),
        ('Scalable ensemble approach',
         'LexGrouping + multi-seed ensembling compress the exponentially large '
         'Fock space into tractable 90-feature vectors.'),
        ('Near-term hardware ready',
         'Implemented on Quandela MerLin — compatible with current photonic '
         'quantum processors without error correction.'),
    ]

    for title, body in conclusions:
        bp = doc.add_paragraph(style='List Bullet')
        add_run(bp, title + ': ', bold=True, size=10.5, color=EPFL_RED)
        add_run(bp, body, size=10.5, color=DARK_SLATE)
        bp.paragraph_format.space_after = Pt(3)

    doc.add_paragraph()

    # ── FOOTER ────────────────────────────────────────────────────────────────
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_shading(footer_p, '1E293B')
    footer_p.paragraph_format.space_before = Pt(6)
    footer_p.paragraph_format.space_after  = Pt(6)
    add_run(footer_p,
            'Team Qedi  ·  EPFL Quantum Hackathon 2026  ·  Quandela Challenge  ·  '
            'github.com/HUmutI/EPFL-Quandela-Qedi',
            size=9, color=RGBColor(0x94, 0xA3, 0xB8), font='Calibri')

    out = '/Users/umut/Desktop/EPFL_ANTI/HPT_QRC_Poster.docx'
    doc.save(out)
    print(f'Saved → {out}')


if __name__ == '__main__':
    build_docx()
